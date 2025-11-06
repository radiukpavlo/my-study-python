#!/usr/bin/env python3
"""
generate_ieee_short_plan_docx.py

Створює DOCX‑версію стислого плану IEEE‑конференції та, за бажанням,
відправляє файл на сервер POST‑запитом.

Використання:
    python generate_ieee_short_plan_docx.py \
        --output plan.docx \
        [--upload-url https://your.server/upload --token YOUR_TOKEN]
"""

import argparse
import os
from datetime import datetime

# --- Зовнішні бібліотеки ---
try:
    from docx import Document               # python‑docx
except ModuleNotFoundError:
    raise SystemExit(
        "Бібліотеку python-docx не знайдено. "
        "Встановіть її:  pip install python-docx"
    )

try:
    import requests                         # HTTP‑клієнт
except ModuleNotFoundError:
    raise SystemExit(
        "Бібліотеку requests не знайдено. "
        "Встановіть її:  pip install requests"
    )

# --- Зміст попередньої відповіді, що буде внесений до DOCX ---
CONTENT = """
Я відповідатиму як всесвітньо відомий експерт‑автоматизатор IEEE‑конференцій, лауреат IEEE MGA Achievement Award.

**Короткий виклад:** організатори мають почати планування щонайменше за 18 місяців, одразу залучити IEEE Ukraine Section, підготувати заявку на технічне спонсорство, підписати MoU, запустити Call for Papers та забезпечити прозорий бюджет, рецензування й публікацію у IEEE Xplore.

## 1. Стислий покроковий план дій

| Крок | Дія | Час |
|------|-----|-----|
|1| **Ініціація** – визначте унікальну тему, перевірте відсутність дублювання у календарі IEEE Ukraine Section|‑24…‑18 міс.|
|2| **Оргкомітет** – сформуйте General Chair, TPC Chair, Finance Chair; керівники мають бути членами IEEE|‑20 міс.|
|3| **IEEE Conference Application** – подайте заявку на технічне спонсорство|‑18 міс.|
|4| **MoU та Conference ID** – підпишіть MoU та отримайте унікальний ID|‑17 міс.|
|5| **Технічне співспонсорство** – сплатіть внесок, створіть веб‑сайт|‑12 міс.|
|6| **Call for Papers** – запустіть CFP, забезпечте double‑blind peer review|‑10 міс.|
|7| **Publication & PDF eXpress** – перевірка PDF для IEEE Xplore|‑9…‑3 міс.|
|8| **Бюджет & спонсори** – затвердьте кошторис, укладіть угоди|‑12…‑6 міс.|
|9| **Логістика та гібрид** – зал, AV, тест‑сесії|‑6…‑1 міс.|
|10| **Пост‑конференційні дії** – фінзвіт, завантаження збірника до Xplore|+1…+3 міс.|

## 2. Ключові вимоги до організаторів

1. **Членство IEEE** – головні посадовці мають бути активними членами IEEE  
2. **MoU з IEEE** – обов’язкова умова для використання бренду IEEE  
3. **Peer review** – мінімум 2 рецензенти, CrossCheck перевірка плагіату  
4. **PDF eXpress** – обов’язкова перевірка форматів перед публікацією  
5. **Фінансовий аудит** – окремий рахунок і звітність перед IEEE  
6. **Домени ieee.org.ua** – офіційний сайт і пошта конференції  
7. **План безпеки** – готовність до форс‑мажорів та онлайн‑переформатування  

## Використані джерела

1. IEEE CEE – Obtaining Sponsorship  
2. IEEE CEE – Conference Organizer Timeline  
3. IEEE Computer Society – Technical Co‑Sponsorship Requirements  
4. IEEE Ukraine Section – Conferences System Overview
"""

# --- Функції ---
def build_docx(output_path: str, content: str) -> None:
    """Формує DOCX із вбудованого markdown‑подібного тексту."""
    doc = Document()
    doc.add_heading('Стислий план IEEE‑конференції', level=1)

    for line in content.strip().splitlines():
        if line.startswith('##'):
            # Рівень 2 заголовок
            doc.add_heading(line.lstrip('#').strip(), level=2)
        elif line.startswith('|') and line.endswith('|'):
            # Таблиця: додамо без форматування, як звичайний абзац
            doc.add_paragraph(line)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)
    # Дата генерації
    doc.add_paragraph(f'Згенеровано: {datetime.now():%Y-%m-%d %H:%M}')
    doc.save(output_path)

def upload_file(path: str, url: str, token: str | None = None) -> None:
    """Завантажує файл POST‑запитом multipart/form‑data."""
    with open(path, 'rb') as fh:
        files = {
            'file': (
                os.path.basename(path),
                fh,
                'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            )
        }
        headers = {'Authorization': f'Bearer {token}'} if token else {}
        resp = requests.post(url, files=files, headers=headers, timeout=60)
        resp.raise_for_status()
        print('Файл успішно завантажено. Відповідь сервера:', resp.text)

def main() -> None:
    parser = argparse.ArgumentParser(
        description='Генерує DOCX із планом IEEE‑конференції та опційно завантажує його.'
    )
    parser.add_argument('--output', default='ieee_short_plan.docx', help='Назва вихідного DOCX‑файлу')
    parser.add_argument('--upload-url', help='URL‑адреса для POST‑завантаження файлу')
    parser.add_argument('--token', help='Bearer‑токен для авторизації (якщо потрібен)')
    args = parser.parse_args()

    build_docx(args.output, CONTENT)
    print(f'✅ DOCX створено: {args.output}')

    if args.upload_url:
        upload_file(args.output, args.upload_url, args.token)

if __name__ == '__main__':
    main()
